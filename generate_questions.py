import os, io, sys, json, textwrap, time
from dotenv import load_dotenv

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt

load_dotenv()
import openai

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MODEL = os.getenv("MODEL", "gpt-4")

if not OPENAI_API_KEY:
    print("ERROR: Set OPENAI_API_KEY in your environment or .env file.")
    sys.exit(1)
openai.api_key = OPENAI_API_KEY

OUT_DIR = "output_files"
os.makedirs(OUT_DIR, exist_ok=True)

def render_latex_to_png(latex, filepath, fontsize=20, dpi=200):
    fig = plt.figure(figsize=(0.01,0.01))
    fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    buf = io.BytesIO()
    plt.axis('off')
    plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', pad_inches=0.1, transparent=True)
    plt.close(fig)
    buf.seek(0)
    with open(filepath, "wb") as f:
        f.write(buf.read())

def create_placeholder_image(text, path, size=(800,400)):
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 18)
    except:
        font = ImageFont.load_default()
    lines = textwrap.wrap(text, width=70)
    y = 20
    for line in lines:
        draw.text((20,y), line, fill="black", font=font)
        y += 22
    img.save(path)

# Read base input
INPUT_DOCX = "sample_input.docx"
base_text = ""
if os.path.exists(INPUT_DOCX):
    d = Document(INPUT_DOCX)
    for p in d.paragraphs:
        base_text += p.text + "\\n"
else:
    base_text = "Base questions not found. Use your own sample_input.docx."

# Prompt GPT to generate JSON output with two questions
prompt = f"""You are an expert assessment writer.
Given the base questions below, create TWO new original multiple-choice math questions similar in style and difficulty.

Constraints:
- Preserve LaTeX in question text. Use $...$ for inline math and $$...$$ for display math.
- If a question needs an image, include "image_filename" and a short description to generate a placeholder image.
- Output valid JSON exactly as specified with an array "questions" of two objects. Example schema for each question:
{{
  "title": "short title",
  "description": "one-line description",
  "question": "question text (LaTeX allowed)",
  "options": ["A...", "B...", "C...", "D...", "E..."],
  "correct_index": 0,
  "explanation": "explain answer (LaTeX allowed)",
  "subject": "Quantitative Math",
  "unit": "Problem Solving",
  "topic": "Numbers and Operations",
  "plusmarks": 1,
  "image_filename": "optional image filename or empty string",
  "image_description": "brief description for placeholder image",
  "latex_images": ["list of LaTeX strings to render separately"]
}}

Base questions:
{base_text}
"""

print("Sending prompt to OpenAI... (this will call the API)")
resp = openai.chat.completions.create(
    model=MODEL,
    messages=[
        {"role":"system","content":"You are a helpful assistant that must output strictly valid JSON and nothing else."},
        {"role":"user","content":prompt}
    ],
    temperature=0.2,
    max_tokens=900
)

content = resp.choices[0].message.content.strip()



# Strip markdown fences if present
if content.startswith("```"):
    parts = content.split("```", 2)
    if len(parts) >= 2:
        content = parts[1].strip()
try:
    data = json.loads(content)
except Exception as e:
    print("Failed to parse JSON from model. Raw content below:\n", content)
    raise e

# Build Word doc from data
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)
doc.add_heading("LLM-Generated Questions (GPT)", level=1)

for idx, q in enumerate(data.get("questions", []), start=1):
    doc.add_paragraph(f"@question {q.get('question')}")
    for opt in q.get('options', []):
        doc.add_paragraph(f"@option {opt}")
    correct = q.get('correct_index', 0)
    correct_text = q.get('options', [])[correct] if q.get('options') else ""
    doc.add_paragraph(f"@@option {correct_text}")
    doc.add_paragraph(f"@explanation {q.get('explanation')}")
    doc.add_paragraph(f"@subject {q.get('subject')}")
    doc.add_paragraph(f"@unit {q.get('unit')}")
    doc.add_paragraph(f"@topic {q.get('topic')}")
    doc.add_paragraph(f"@plusmarks {q.get('plusmarks')}")
    # create placeholder image if described
    img_fn = q.get('image_filename','').strip()
    img_desc = q.get('image_description','').strip()
    if img_fn and img_desc:
        img_path = os.path.join(OUT_DIR, img_fn)
        create_placeholder_image(img_desc, img_path)
        doc.add_picture(img_path, width=Inches(5))
    # render latex images
    for i, latex in enumerate(q.get('latex_images', [])):
        outimg = os.path.join(OUT_DIR, f"q{idx}_latex_{i+1}.png")
        try:
            render_latex_to_png(latex, outimg)
            doc.add_picture(outimg, width=Inches(5))
        except Exception as e:
            print("LaTeX render failed:", e)

out_docx = os.path.join(OUT_DIR, "llm_generated_questions.docx")
doc.save(out_docx)
print("Saved:", out_docx)
print("Done.")



try:
    import dotenv
    print("dotenv imported")
except ImportError:
    print("dotenv NOT found")

try:
    import docx
    print("docx imported")
except ImportError:
    print("docx NOT found")
